import docx
import re
import spacy
from typing import List, Dict, Optional
import os
import sys
import json

# Load spaCy for sentence splitting
try:
    nlp = spacy.load("en_core_web_sm")
except OSError:
    print("Install spaCy model:")
    print("  python -m spacy download en_core_web_sm")
    sys.exit(1)

# Defaults (overridden by parse_docx args)
MAX_TOKENS = 500
MIN_CHUNK_TOKENS = 20

def count_tokens(text: str) -> int:
    if not text or not text.strip():
        return 0
    return max(1, len(text.strip()) // 4)

def split_into_sentences(text: str) -> List[str]:
    doc = nlp(text)
    return [s.text.strip() for s in doc.sents if s.text.strip()]

def split_long_text(text: str, max_tokens: int = MAX_TOKENS) -> List[str]:
    if not text.strip():
        return []
    sents = split_into_sentences(text)
    if not sents:
        return [text.strip()]
    chunks, cur = [], ""
    for sent in sents:
        cand = cur + (" " + sent if cur else sent)
        if count_tokens(cand) > max_tokens and cur:
            chunks.append(cur.strip())
            cur = sent
        else:
            cur = cand
    if cur.strip():
        chunks.append(cur.strip())
    return chunks

def is_heading(text: str) -> bool:
    return bool(text and re.match(r"^\d+(\.\d+)*\s", text.strip()))

def extract_section_number(text: str) -> Optional[str]:
    m = re.match(r"^(\d+(\.\d+)*)", text.strip())
    return m.group(1) if m else None

def extract_title(text: str) -> str:
    return re.sub(r"^\d+(\.\d+)*\s*", "", text.strip())

def merge_small_chunks(chunks: List[Dict]) -> List[Dict]:
    merged, i = [], 0
    while i < len(chunks):
        curr = chunks[i].copy()
        if curr["chunk_type"] == "heading":
            merged.append(curr); i += 1; continue
        while (
            i + 1 < len(chunks)
            and curr["tokens"] < MIN_CHUNK_TOKENS
            and chunks[i+1]["chunk_type"] != "heading"
            and curr["section_id"] == chunks[i+1]["section_id"]
        ):
            nxt = chunks[i+1]
            combined = curr["content"] + " " + nxt["content"]
            if count_tokens(combined) <= MAX_TOKENS:
                curr["content"] = combined
                curr["tokens"] = count_tokens(combined)
                i += 1
            else:
                break
        merged.append(curr); i += 1
    return merged

def get_first_page_paragraphs(doc) -> List[str]:
    """
    Extract paragraphs and table content from the first page only.
    This is a simplified approach - in reality, determining exact page breaks
    in python-docx is complex, so we'll use a heuristic.
    """
    first_page_content = []
    
    # Method 1: Look for explicit page breaks or section breaks
    for para in doc.paragraphs:
        # Check if this paragraph has a page break before it
        if para._element.xpath('.//w:br[@w:type="page"]'):
            break
            
        text = para.text.strip()
        if text:
            first_page_content.append(text)
            
        # Stop at common first-page endings
        if (text.lower().startswith("contents") or 
            text.lower().startswith("table of contents") or
            is_heading(text) and extract_section_number(text) == "1"):
            break
    
    # Add table content from first page (before any heading appears)
    seen_heading = False
    for table in doc.tables:
        # Stop processing tables once we see a numbered heading
        if seen_heading:
            break
            
        for row in table.rows:
            for cell in row.cells:
                cell_text = cell.text.strip()
                if cell_text:
                    first_page_content.append(cell_text)
                    # Check if this cell contains a heading - if so, we've moved past first page
                    if is_heading(cell_text) and extract_section_number(cell_text) == "1":
                        seen_heading = True
                        break
            if seen_heading:
                break
    
    # If we didn't find any clear break, take first ~15 paragraphs as heuristic
    if len(first_page_content) == 0:
        first_page_content = [p.text.strip() for p in doc.paragraphs[:15] if p.text.strip()]
    
    return first_page_content

def parse_docx(
    file_path: str,
    max_tokens: int = 500,
    min_chunk_tokens: int = 20
) -> List[Dict]:
    """
    1) Parse and chunk ONLY the first page as preamble
    2) Then resume normal heading-based parsing for the rest
    """
    global MAX_TOKENS, MIN_CHUNK_TOKENS
    MAX_TOKENS, MIN_CHUNK_TOKENS = max_tokens, min_chunk_tokens

    if not os.path.exists(file_path):
        raise FileNotFoundError(f"{file_path} not found")
    doc = docx.Document(file_path)

    # --- Setup preamble section ---
    chunks: List[Dict] = [{
        "section_id": "0_preamble",
        "parent_section": None,
        "title": "Preamble",
        "chunk_id": "0_preamble_0",
        "content": "",
        "chunk_type": "heading",
        "position": 0,
        "tokens": 0
    }]
    current_section = "0_preamble"
    parent = None
    pos = 0

    def add_text_to_preamble(text: str):
        nonlocal pos
        for sub in split_long_text(text, MAX_TOKENS):
            pos += 1
            chunks.append({
                "section_id": current_section,
                "parent_section": parent,
                "title": None,
                "chunk_id": f"{current_section}_{pos}",
                "content": sub,
                "chunk_type": "paragraph",
                "position": pos,
                "tokens": count_tokens(sub)
            })

    def add_text_to_current_section(text: str):
        nonlocal pos
        for sub in split_long_text(text, MAX_TOKENS):
            pos += 1
            chunks.append({
                "section_id": current_section,
                "parent_section": parent,
                "title": None,
                "chunk_id": f"{current_section}_{pos}",
                "content": sub,
                "chunk_type": "paragraph",
                "position": pos,
                "tokens": count_tokens(sub)
            })

    # --- 1) FIRST PAGE ONLY as preamble ---
    first_page_texts = get_first_page_paragraphs(doc)
    first_page_end_idx = len(first_page_texts)
    
    # Add first page content to preamble
    for text in first_page_texts:
        add_text_to_preamble(text)

    # --- 2) Find where to resume normal parsing ---
    resume_idx = 0
    para_count = 0
    
    for idx, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if not text:
            continue
            
        para_count += 1
        if para_count > len(first_page_texts):
            resume_idx = idx
            break

    # --- 3) NORMAL PARSING from resume point onward ---
    seen_heading = False
    
    for para in doc.paragraphs[resume_idx:]:
        txt = para.text.strip()
        if not txt:
            continue

        if is_heading(txt):
            seen_heading = True
            sec = extract_section_number(txt)
            if sec:
                current_section = sec
                parent = ".".join(sec.split(".")[:-1]) if "." in sec else None
                pos = 0
                title = extract_title(txt)
                chunks.append({
                    "section_id": current_section,
                    "parent_section": parent,
                    "title": title,
                    "chunk_id": f"{current_section}_0",
                    "content": title,
                    "chunk_type": "heading",
                    "position": 0,
                    "tokens": count_tokens(title)
                })
            continue

        # Only add content if we're in a proper section (past preamble)
        if seen_heading or current_section != "0_preamble":
            add_text_to_current_section(txt)

    # --- 4) Add header content to preamble (optional) ---
    if doc.sections:
        hdr = doc.sections[0].header
        current_section = "0_preamble"  # Reset to preamble for header content
        for para in hdr.paragraphs:
            if para.text.strip():
                add_text_to_preamble(para.text.strip())

    # --- Finalize: merge tiny bits & reassign IDs ---
    merged = merge_small_chunks(chunks)
    counters: Dict[str,int] = {}
    for c in merged:
        sec = c["section_id"]
        idx = counters.get(sec, 0)
        c["chunk_id"] = f"{sec}_{idx}"
        c["position"] = idx
        counters[sec] = idx + 1

    return merged

def save_as_json(chunks: List[Dict], out_path: str):
    os.makedirs(os.path.dirname(out_path), exist_ok=True)
    with open(out_path, "w", encoding="utf-8") as f:
        json.dump(chunks, f, indent=2, ensure_ascii=False)