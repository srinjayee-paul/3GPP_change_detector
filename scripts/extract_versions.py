#!/usr/bin/env python3
# scripts/extract_versions.py

import re
import json
import argparse
from docx import Document

def get_first_page_paragraphs(doc: Document, max_paragraphs=20):
    """
    Get paragraphs from approximately the first page.
    Since python-docx doesn't have direct page breaks, we'll take the first N paragraphs.
    """
    first_page_paras = []
    para_count = 0
    
    for para in doc.paragraphs:
        if para_count >= max_paragraphs:
            break
        first_page_paras.append(para)
        para_count += 1
    
    return first_page_paras

def extract_from_paragraphs(doc: Document):
    """
    Search only first page paragraphs for a line containing '3GPP TS' and a version.
    Returns the first matching full paragraph text, or None.
    """
    # More flexible pattern to catch variations
    pattern = re.compile(r"3GPP\s+TS.*?V?\d+\.\d+\.\d+", re.IGNORECASE)
    
    # Only search first page paragraphs
    first_page_paras = get_first_page_paragraphs(doc)
    
    for para in first_page_paras:
        text = para.text.strip()
        if not text:
            continue
        
        if pattern.search(text):
            return text
    return None

def extract_from_tables(doc: Document):
    """
    Search only first page tables for a cell containing '3GPP TS' + version.
    Returns the first matching full cell text, or None.
    """
    pattern = re.compile(r"3GPP\s+TS.*?V?\d+\.\d+\.\d+", re.IGNORECASE)
    
    # Get first few tables (likely on first page)
    table_count = 0
    max_tables = 5
    
    for tbl in doc.tables:
        if table_count >= max_tables:
            break
        table_count += 1
        
        for row in tbl.rows:
            for cell in row.cells:
                text = cell.text.strip().replace("\n", " ")
                if pattern.search(text):
                    return text
    return None

def extract_document_properties(doc: Document):
    """
    Try to extract version info from document properties/metadata.
    """
    try:
        # Check document title
        if hasattr(doc.core_properties, 'title') and doc.core_properties.title:
            title = doc.core_properties.title
            pattern = re.compile(r"3GPP\s+TS\s+\d+\.\d+\s+V\d+\.\d+\.\d+", re.IGNORECASE)
            if pattern.search(title):
                return title
                
        # Check document subject
        if hasattr(doc.core_properties, 'subject') and doc.core_properties.subject:
            subject = doc.core_properties.subject
            if pattern.search(subject):
                return subject
    except:
        pass
    return None

def extract_from_headers_footers(doc: Document):
    """
    Search headers and footers for version information.
    """
    pattern = re.compile(r"3GPP\s+TS\s+\d+\.\d+\s+V\d+\.\d+\.\d+", re.IGNORECASE)
    
    for section in doc.sections:
        # Check headers
        if section.header:
            for para in section.header.paragraphs:
                text = para.text.strip()
                if pattern.search(text):
                    return text
        
        # Check footers  
        if section.footer:
            for para in section.footer.paragraphs:
                text = para.text.strip()
                if pattern.search(text):
                    return text
    return None

def extract_alternative_patterns(doc: Document):
    """
    Look for alternative version patterns that might be in the second file.
    """
    patterns = [
        # Look for TS followed by version without "3GPP" prefix
        re.compile(r"TS\s+\d+\.\d+\s+V\d+\.\d+\.\d+", re.IGNORECASE),
        # Look for version in title-like format
        re.compile(r"V\d+\.\d+\.\d+\s*\(\d{4}-\d{2}\)", re.IGNORECASE),
        # Look for version with release info
        re.compile(r"Version\s+\d+\.\d+\.\d+.*Release\s+\d+", re.IGNORECASE)
    ]
    
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text or len(text) > 200:  # Skip long paragraphs
            continue
            
        for pattern in patterns:
            if pattern.search(text):
                return text
    
    return None

def extract_release_info(doc: Document, start_text: str):
    """
    Extract release information from version string or nearby content.
    Returns the release info in format "(Release X)" or None.
    """
    if not start_text:
        return None
    
    # First, try to extract release from the version string itself
    # Look for patterns like "V17.12.0" where the major version (17) indicates Release 17
    version_pattern = re.compile(r"V(\d+)\.\d+\.\d+", re.IGNORECASE)
    version_match = version_pattern.search(start_text)
    
    if version_match:
        major_version = int(version_match.group(1))
        return f"(Release {major_version})"
    
    # Fallback: look for explicit "Release X" text
    release_pattern = re.compile(r"release\s*(\d+)", re.IGNORECASE)
    
    # Check if release info is in the same line as version
    if release_pattern.search(start_text):
        match = release_pattern.search(start_text)
        return f"(Release {match.group(1)})"
    
    # Search first page paragraphs for release info
    first_page_paras = get_first_page_paragraphs(doc)
    
    collecting = False
    for para in first_page_paras:
        text = para.text.strip()
        if not text:
            continue
        if collecting:
            if release_pattern.search(text):
                match = release_pattern.search(text)
                return f"(Release {match.group(1)})"
        if start_text and text == start_text:
            collecting = True
    
    return None

def extract_version(path: str):
    """Enhanced version extraction focused on first page content."""
    doc = Document(path)
    
    # Try extraction methods in order of preference, but limit to first page
    extraction_methods = [
        ("document_properties", lambda: extract_document_properties(doc)),
        ("headers_footers", lambda: extract_from_headers_footers(doc)),
        ("paragraphs", lambda: extract_from_paragraphs(doc)),
        ("tables", lambda: extract_from_tables(doc)),
    ]
    
    ver_line = None
    
    for method_name, method_func in extraction_methods:
        try:
            result = method_func()
            if result:
                ver_line = result
                break
        except Exception as e:
            print(f"Warning: {method_name} extraction failed: {e}")
            continue
    
    # Try to extract release info
    rel_line = None
    if ver_line:
        rel_line = extract_release_info(doc, ver_line)
    
    return {
        "version_line": ver_line,
        "release_info": rel_line
    }

def main():
    parser = argparse.ArgumentParser(
        description="Extract 3GPP spec version & release info from two DOCX files"
    )
    parser.add_argument("--old", required=True, help="path to old-release DOCX")
    parser.add_argument("--new", required=True, help="path to new-release DOCX")
    parser.add_argument("--out", default="data/processed/versions.json",
                        help="output JSON path")
    parser.add_argument("--debug", action="store_true", 
                        help="show debug information")
    args = parser.parse_args()

    print("Extracting from old file...")
    old_info = extract_version(args.old)
    
    print("Extracting from new file...")
    new_info = extract_version(args.new)

    if args.debug:
        print(f"Old file method: {old_info.get('extraction_method')}")
        print(f"New file method: {new_info.get('extraction_method')}")

    result = {
        "rel_old": old_info,
        "rel_new": new_info
    }

    # ensure output dir exists
    import os
    os.makedirs(os.path.dirname(args.out), exist_ok=True)

    with open(args.out, "w", encoding="utf-8") as f:
        json.dump(result, f, indent=2, ensure_ascii=False)

    print(f"✓ Extracted versions written to {args.out}")
    print(json.dumps(result, indent=2, ensure_ascii=False))

if __name__ == "__main__":
    main()